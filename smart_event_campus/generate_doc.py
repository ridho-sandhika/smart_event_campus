from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Dokumentasi Sistem Manajemen Event', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle = doc.add_paragraph('Ujian Akhir Semester (UAS)')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('1. Pendahuluan', level=1)
doc.add_paragraph('Aplikasi ini adalah Sistem Manajemen Event berbasis Web yang memungkinkan pengguna untuk melakukan pendaftaran pada suatu event. Terdapat dua role utama dalam sistem ini:')
ul = doc.add_paragraph(style='List Bullet')
ul.add_run('Pengguna Umum (User): ').bold = True
ul.add_run('Dapat melihat daftar event yang tersedia dan mendaftarkan diri pada event tersebut.')
ul2 = doc.add_paragraph(style='List Bullet')
ul2.add_run('Administrator (Admin): ').bold = True
ul2.add_run('Memiliki hak akses penuh untuk mengelola (CRUD) data event, melihat daftar pendaftar, serta mengunduh laporan pendaftaran dalam bentuk PDF.')

doc.add_heading('2. Struktur Database', level=1)
doc.add_paragraph('Sistem ini menggunakan database MySQL. Terdapat beberapa tabel utama yang digunakan (berdasarkan database.sql):')
doc.add_paragraph('1. Tabel Users: Menyimpan data admin dan user.', style='List Number')
doc.add_paragraph('2. Tabel Events: Menyimpan informasi terkait acara yang tersedia.', style='List Number')
doc.add_paragraph('3. Tabel Registrations: Menyimpan data peserta yang mendaftar.', style='List Number')

doc.add_heading('3. Penjelasan Fitur dan Antarmuka (Tampilan)', level=1)

# Fitur 1
doc.add_heading('3.1. Halaman Login (login.php)', level=2)
doc.add_paragraph('Fungsi: Memvalidasi kredensial pengguna (username dan password).', style='List Bullet')
doc.add_paragraph('Tampilan: Form login "Secure Login" yang rapi dengan validasi input.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO LOGIN DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Fitur 2
doc.add_heading('3.2. Halaman Utama / Beranda (index.php)', level=2)
doc.add_paragraph('Fungsi: Menampilkan daftar event yang tersedia.', style='List Bullet')
doc.add_paragraph('Tampilan: Terdapat daftar event (seperti lomba, seminar) dalam bentuk card.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO BERANDA DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Fitur 3
doc.add_heading('3.3. Halaman Dashboard Admin (admin.php)', level=2)
doc.add_paragraph('Fungsi: Menampilkan seluruh statistik event dalam bentuk visual.', style='List Bullet')
doc.add_paragraph('Tampilan: Ringkasan data kegiatan dan grafik distribusi.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO ADMIN ATAS DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO ADMIN BAWAH DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Fitur 4
doc.add_heading('3.4. Halaman Kelola Event (create_event.php & edit_event.php)', level=2)
doc.add_paragraph('Fungsi: Memasukkan rincian event baru atau mengubah yang sudah ada.', style='List Bullet')
doc.add_paragraph('Tampilan: Form input data event yang lengkap.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO TAMBAH EVENT DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO EDIT EVENT DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Fitur 5
doc.add_heading('3.5. Halaman Pendaftaran Event (register_event.php)', level=2)
doc.add_paragraph('Fungsi: Mengisi form pendaftaran acara.', style='List Bullet')
doc.add_paragraph('Tampilan: Form pendaftaran yang interaktif.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO PENDAFTARAN DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Fitur 6
doc.add_heading('3.6. Laporan dan Ekspor PDF', level=2)
doc.add_paragraph('Fungsi: Menghasilkan dokumen cetak berisi daftar seluruh kegiatan atau pendaftar.', style='List Bullet')
doc.add_paragraph('Tampilan: Halaman siap cetak.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO CETAK LAPORAN KEGIATAN DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO CETAK LAPORAN PENDAFTAR DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Section 4
doc.add_heading('4. Fitur-Fitur Tambahan & Pembaruan (Patch Terbaru)', level=1)

ul_patch1 = doc.add_paragraph(style='List Bullet')
ul_patch1.add_run('Desain Premium Glassmorphism (Tema Indigo-Violet): ').bold = True
ul_patch1.add_run('Warna utama dirombak menjadi skema indigo-violet premium dengan estetika transparan (glassmorphism) yang bersih dan mewah.')

ul_patch2 = doc.add_paragraph(style='List Bullet')
ul_patch2.add_run('Mode Malam / Siang (Dark & Light Mode Toggle): ').bold = True
ul_patch2.add_run('Kini tersedia opsi untuk mengubah antarmuka menjadi mode gelap atau terang yang akan merubah palet warna halaman secara keseluruhan, sehingga memberikan pengalaman visual yang lebih baik.')

ul_patch3 = doc.add_paragraph(style='List Bullet')
ul_patch3.add_run('Navbar Garis Tiga (Hamburger Overlay Menu): ').bold = True
ul_patch3.add_run('Seluruh navigasi, termasuk Admin Login dan Jadwal Sholat, disederhanakan ke dalam menu hamburger di semua resolusi perangkat (HP maupun Laptop). Tampilannya dirombak dengan sistem overlay yang interaktif dan responsif, terbebas dari bug layout.')

ul_patch4 = doc.add_paragraph(style='List Bullet')
ul_patch4.add_run('Icon Mengambang (Floating Action Buttons): ').bold = True
ul_patch4.add_run('Pengumuman penting kini ditampilkan lewat sebuah icon (badge) mengambang di pojok layar utama, serupa dengan icon Customer Service. Saat di-klik, pengumuman akan tampil elegan dalam mode layar penuh (fullscreen overlay).')

ul_patch5 = doc.add_paragraph(style='List Bullet')
ul_patch5.add_run('Perbaikan Tombol Kembali (Back Button): ').bold = True
ul_patch5.add_run('Seluruh tombol kembali pada halaman-halaman turunan (Admin, Pendaftaran, Detail, dll.) dipindahkan secara rapi dan independen dari struktur navbar, meningkatkan fleksibilitas dan kenyamanan navigasi pengguna.')

doc.add_heading('5. Kesimpulan', level=1)
doc.add_paragraph('Sistem ini dibangun dengan arsitektur PHP native yang terstruktur, memudahkan pengelolaan event mulai dari publikasi hingga pencatatan peserta pendaftaran dengan antarmuka yang sangat modern (Glassmorphism, Dark Mode) dan responsif penuh di semua perangkat.')

doc.save('Dokumentasi_Project_Update.docx')
print("Word document generated as Dokumentasi_Project_Update.docx.")
